import logging
import tempfile
from pathlib import Path

import docx

from parser.doc_converter import convert_doc_to_docx

logger = logging.getLogger(__name__)


def _parse_docx(file_path: Path) -> str:
    doc = docx.Document(str(file_path))
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text.append(row_text)
    return "\n".join(text)


def parse_word(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".doc":
        docx_path = convert_doc_to_docx(str(path))
        try:
            return _parse_docx(docx_path)
        finally:
            docx_path.unlink(missing_ok=True)

    if ext == ".docx":
        return _parse_docx(path)

    raise ValueError(f"Unsupported Word format: {ext}")
