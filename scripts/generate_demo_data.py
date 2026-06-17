"""Generate demo catalog and input files for the parser."""

from pathlib import Path

import pandas as pd
from docx import Document

BASE = Path(__file__).resolve().parent.parent
CATALOG_PATH = BASE / "data" / "catalog" / "catalog.xlsx"
INPUT_PATH = BASE / "data" / "input" / "input.docx"

SHEETS = {
    "Дорожные": [
        {
            "Наименование": "Светодиодный светильник ДКУ-150",
            "dimensions": "600x600 мм",
            "Характеристики": "Мощность: 150 Вт; Цветовая температура: 4000 К; Световой поток: 18000 Лм",
        },
        {
            "Наименование": "Светодиодный светильник ДКУ-100",
            "dimensions": "500x500 мм",
            "Характеристики": "Мощность: 100 Вт; Цветовая температура: 5000 К; Световой поток: 12000 Лм",
        },
    ],
    "Светильники на опоре": [
        {
            "Наименование": "Консольный светильник КС-120",
            "dimensions": "800x300 мм",
            "Характеристики": "Мощность: 120 Вт; Цветовая температура: 4000 К; Световой поток: 15000 Лм",
        },
        {
            "Наименование": "Консольный светильник КС-90",
            "dimensions": "700x250 мм",
            "Характеристики": "Мощность: 90 Вт; Цветовая температура: 3000 К; Световой поток: 11000 Лм",
        },
    ],
    "Осветительные комплексы": [
        {
            "Наименование": "Комплекс ОК-200",
            "dimensions": "1200x600 мм",
            "Характеристики": "Мощность: 200 Вт; Цветовая температура: 4000 К; Световой поток: 24000 Лм",
        },
    ],
    "Промышленные": [
        {
            "Наименование": "Промышленный светильник ПС-80",
            "dimensions": "1200x150 мм",
            "Характеристики": "Мощность: 80 Вт; Цветовая температура: 5000 К; Световой поток: 9000 Лм",
        },
    ],
    "Аварийные": [
        {
            "Наименование": "Аварийный светильник АС-20",
            "dimensions": "300x100 мм",
            "Характеристики": "Мощность: 20 Вт; Цветовая температура: 6500 К; Световой поток: 2000 Лм",
        },
    ],
    "Прожекторы": [
        {
            "Наименование": "LED прожектор ПР-50",
            "dimensions": "250x200 мм",
            "Характеристики": "Мощность: 50 Вт; Цветовая температура: 5000 К; Световой поток: 6000 Лм",
        },
    ],
}


def create_catalog():
    CATALOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(CATALOG_PATH, engine="openpyxl") as writer:
        for sheet_name, rows in SHEETS.items():
            pd.DataFrame(rows).to_excel(writer, sheet_name=sheet_name, index=False)
    print(f"Catalog created: {CATALOG_PATH}")


def create_input_docx():
    INPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Спецификация светотехнического оборудования", level=1)
    doc.add_paragraph(
        "Заказчик: ООО «Городское освещение». "
        "Поставка светодиодного оборудования для реконструкции уличного освещения."
    )

    doc.add_heading("Позиции закупки", level=2)
    doc.add_paragraph(
        "1. Светодиодный светильник для дорожного освещения — 25 шт. "
        "Размеры: 600x600 мм. Мощность: 150 Вт. Цветовая температура: 4000 К. "
        "Световой поток: 18000 Лм."
    )
    doc.add_paragraph(
        "2. Консольный светильник на опору — 15 шт. "
        "Размеры: 800x300 мм. Мощность: 120 Вт. Цветовая температура: 4000 К. "
        "Световой поток: 15000 Лм."
    )
    doc.add_paragraph(
        "3. LED прожектор для архитектурной подсветки — 8 шт. "
        "Размеры: 250x200 мм. Мощность: 50 Вт. Цветовая температура: 5000 К. "
        "Световой поток: 6000 Лм."
    )

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["Наименование", "Кол-во", "Мощность", "Размеры"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    table.rows[1].cells[0].text = "Аварийный светильник"
    table.rows[1].cells[1].text = "10"
    table.rows[1].cells[2].text = "20 Вт"
    table.rows[1].cells[3].text = "300x100 мм"

    doc.save(INPUT_PATH)
    print(f"Input document created: {INPUT_PATH}")


if __name__ == "__main__":
    create_catalog()
    create_input_docx()
